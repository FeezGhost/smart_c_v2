from base64 import b64decode
from PyPDF2 import PdfWriter
from docx import Document
from io import BytesIO
import requests
import os
from django.conf import settings
import re


def create_empty_pdf(file_path):
    pdf_writer = PdfWriter()
    with open(file_path, 'wb') as output_pdf:
        pdf_writer.write(output_pdf)

def convert_latex_format(latex_code):
    # Remove newlines and extra spaces
    latex_code = re.sub(r"\n\s*", "", latex_code)

    # Remove unnecessary commands
    # latex_code = re.sub(r"\\usepackage\{.*?\}", "", latex_code)
    # latex_code = re.sub(r"\\geometry\{.*?\}", "", latex_code)
    # latex_code = re.sub(r"\\subsection\*", r"\\section*", latex_code)
    # latex_code = re.sub(r"\\\\",r"\\", latex_code)
    latex_code = re.sub(r"\\\\",r"\\", latex_code)
    # latex_code = latex_code.encode('utf-8').decode('unicode-escape')

    return latex_code

# Compiles LaTeX code into PDF and returns the PDF content as bytes
def compile_latex_to_pdf(latex_code, filename ):
    
    with open(filename, 'r') as file:
        content = file.read()
    
    url = 'https://texlive.net/cgi-bin/latexcgi'

    # Prepare the form data
    form_data = {
        'return': 'pdf',
        'engine': 'pdflatex',
        'filecontents[]': content.encode('utf-8'),
        'filename[]': 'document.tex'
    }

    # Send the POST request
    response = requests.post(url, files=form_data)

    # Send the POST request
    # response = requests.post(url, headers=headers, files=form_data)
    print(response.content)
    if response.status_code == 200:
        # Remove the file extension from the text file name
        pdf_file_path = filename.rsplit('.', 1)[0] + '.pdf'

        with open(pdf_file_path, 'wb') as file:
            file.write(response.content)
        return pdf_file_path

    else:
        print("PDF compilation failed. Status code:", response.status_code)
        return None

# Function to upload a file to file.io and return the download link
def upload_to_file_io(file_path):
  with open(file_path, "rb") as f:
    response = requests.post("https://file.io", files={"file": f})
    if response.status_code == 200:
      return response.json()["link"]
    else:
      raise Exception(f"Failed to upload file to file.io: {response.text}")

# Function to save a text file with the given content and file name
def save_text_file(content, file_name, folder="file_generation/tex_templates/"):
    file_path = os.path.join(folder, f"{file_name}")
    with open(file_path, "w") as f:
        f.write(content)
    full_path = os.path.join(settings.BASE_DIR, file_path)
    full_path = full_path.replace("\\", "/")
    return full_path


# Function to generate a Word document with the given content and document type
def generate_word_document(content, document_type):
  document = Document()

  document.add_heading(f'{document_type.capitalize()}', 0)
  document.add_paragraph(content)

  output_buffer = BytesIO()
  document.save(output_buffer)
  output_buffer.seek(0)

  return output_buffer

